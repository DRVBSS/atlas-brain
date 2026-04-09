"""Word document extraction via python-docx."""

from pathlib import Path

from atlas_brain.models import ProcessedDocument


def extract(file_path: Path) -> ProcessedDocument:
    """Extract text from DOCX files."""
    from docx import Document

    doc = Document(str(file_path))

    paragraphs = []
    sections = []
    title = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect headings
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            sections.append(text)
            level = para.style.name.replace("Heading", "").strip()
            try:
                level_num = int(level)
            except ValueError:
                level_num = 1
            prefix = "#" * level_num
            paragraphs.append(f"{prefix} {text}")

            if title is None:
                title = text
        else:
            paragraphs.append(text)

    # Extract tables
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(" | ".join(cells))
        if table_rows:
            paragraphs.append("\n".join(table_rows))

    full_text = "\n\n".join(paragraphs)

    if not title:
        title = file_path.stem

    # Try to get author from core properties
    author = None
    try:
        if doc.core_properties and doc.core_properties.author:
            author = doc.core_properties.author
    except Exception:
        pass

    created_date = None
    try:
        if doc.core_properties and doc.core_properties.created:
            created_date = doc.core_properties.created.isoformat()
    except Exception:
        pass

    return ProcessedDocument(
        text=full_text,
        title=title,
        author=author,
        created_date=created_date,
        word_count=len(full_text.split()),
        sections=sections,
        metadata={},
    )
